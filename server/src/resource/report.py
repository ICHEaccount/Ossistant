import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.style import WD_STYLE

from datetime import datetime
from io import BytesIO

from db_conn.neo4j.models import *


def add_title(doc, name):
    para_relation = doc.add_paragraph()
    para_run = para_relation.add_run(name)
    para_run.bold = True
    para_run.underline = True
    para_run.font.size = Pt(20)


def add_img(case, doc, target):
    image_path = f'./docs/report/{case.case_id}/{target}.png'

    if not target == 'relation':
        if not target == 'whole':
            doc.add_paragraph()
        sub_para = doc.add_paragraph()
        sub_run = sub_para.add_run(target)
        sub_run.bold = True
        sub_run.size = Pt(15)

    img_para = doc.add_paragraph()
    img_run = img_para.add_run()
    img_run.add_picture(image_path, width=Cm(16))
    img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def report(case):
    io_stream = BytesIO()
    doc = Document()

    # Add first section
    section = doc.sections[0] if doc.sections else doc.add_section()

    # Set margin
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(1.91)
    section.top_margin = Cm(1.91)
    section.bottom_margin = Cm(2.54)

    # Set header
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = f"OSSISTANT Report\t\t{datetime.today().strftime('%Y%m%d')}"

    # Set footer
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    # Set page number
    page_number = footer_paragraph.add_run()
    page_number.font.size = Pt(10)
    # page_number.font.name = "Arial"
    page_number.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    xml = '<w:fldSimple %s w:instr="PAGE" />' % nsdecls('w')
    r_element = parse_xml(xml)
    footer_paragraph._p.insert_element_before(r_element, 'w:pPr')

    '''
    First page (cover page)
    '''
    # Add title
    for i in range(3):
        doc.add_paragraph()
    title = "OSSISTANT Report"
    doc.add_heading(title, 0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for i in range(15):
        doc.add_paragraph()

    # Add table
    table_case_info = doc.add_table(rows=4, cols=2)
    table_case_info.style = doc.styles['Table Grid']
    table_case_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    row_case_name = table_case_info.rows[0].cells
    row_case_name[0].text = 'Case Name'
    row_case_name[1].text = case.case_name

    row_case_number = table_case_info.rows[1].cells
    row_case_number[0].text = 'Case Number'
    row_case_number[1].text = case.case_num

    row_investigator = table_case_info.rows[2].cells
    row_investigator[0].text = 'Investigator'
    row_investigator[1].text = case.investigator

    row_desc = table_case_info.rows[3].cells
    row_desc[0].text = 'Description'
    row_desc[1].text = case.description

    # Table alignment
    for row in table_case_info.rows:
        for idx, cell in enumerate(row.cells):
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            if idx % 2 == 0:  # Left column
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:  # Right column
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Table width set
    for column in table_case_info.columns:
        for cell in column.cells:
            cell.width = Pt(100)

    doc.add_page_break()  # End of cover page

    '''
    Start of main contents
    '''
    # 1-1. Relation
    add_title(doc, 'Relation')
    add_img(case, doc, 'relation')

    # 1-2. Timeline
    add_title(doc, 'Timeline')
    add_img(case, doc, 'whole')
    add_img(case, doc, 'suspect')
    add_img(case, doc, 'domain')

    for key, node_obj in NODE_LIST.items():
        get_node_flag, node_data = node_obj.get_all_nodes_list(case_id=case.case_id, is_export=True)
        if get_node_flag is True:
            node_df = pd.DataFrame(node_data)

            # Table
            if not node_df.empty:
                for i in range(2):
                    doc.add_paragraph()
                add_title(doc, name=key)
                node_table = doc.add_table(rows=node_df.shape[0]+1, cols=node_df.shape[1])
                node_table.style = doc.styles['Table Grid']

                for col_num, col_name in enumerate(node_df.columns):  # Add header
                    cell = node_table.cell(0, col_num)
                    # if key == 'Post' and col_num == 0: # not working
                    #     cell.width = Pt(13)
                    cell.text = col_name
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

                    shading_elm = parse_xml(r'<w:shd {} w:fill="0e0f37"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading_elm)

                for row_num, (_, row) in enumerate(node_df.iterrows(), start=1):
                    for col_num, value in enumerate(row):
                        cell = node_table.cell(row_num, col_num)
                        cell.text = str(value)
                        cell.paragraphs[0].runs[0].font.size = Pt(8)

    # Final
    doc.save(io_stream)
    return io_stream
